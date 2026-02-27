"""
Export the finished report to Markdown, DOCX, and/or PDF.
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from tdrcreator.security.logger import get_logger

_log = get_logger("report.exporter")


def export_markdown(markdown: str, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(markdown, encoding="utf-8")
    _log.metric("export_md", path=str(output_path), chars=len(markdown))


def export_docx(markdown: str, output_path: Path) -> None:
    """
    Convert markdown to DOCX via python-docx (simple approach).
    For richer conversion, pandoc is preferred when available.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Try pandoc first (best quality)
    if _pandoc_available():
        _pandoc_convert(markdown, output_path, to="docx")
        return

    # Fallback: python-docx
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise RuntimeError(
            "python-docx not installed (pip install python-docx) "
            "and pandoc not available in PATH."
        )

    doc = Document()
    for line in markdown.split("\n"):
        line = line.rstrip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("---"):
            doc.add_paragraph("_" * 40)
        elif line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")

    doc.save(str(output_path))
    _log.metric("export_docx", path=str(output_path))


def export_pdf(markdown: str, output_path: Path) -> None:
    """
    Convert markdown to PDF via pandoc (requires pandoc + xelatex or wkhtmltopdf).
    Falls back to reportlab if pandoc is unavailable.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if _pandoc_available():
        _pandoc_convert(markdown, output_path, to="pdf")
        return

    # Fallback: reportlab (basic)
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
    except ImportError:
        raise RuntimeError(
            "Neither pandoc nor reportlab is installed. "
            "Install pandoc (https://pandoc.org) or: pip install reportlab"
        )

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(output_path), pagesize=A4)
    story = []

    for line in markdown.split("\n"):
        line = line.rstrip()
        if line.startswith("# "):
            story.append(Paragraph(line[2:], styles["Title"]))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], styles["Heading2"]))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:], styles["Heading3"]))
        elif line.startswith("---"):
            story.append(Spacer(1, 0.2 * cm))
        elif line:
            safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe_line, styles["Normal"]))
        else:
            story.append(Spacer(1, 0.3 * cm))

    doc.build(story)
    _log.metric("export_pdf", path=str(output_path))


def _pandoc_available() -> bool:
    try:
        result = subprocess.run(
            ["pandoc", "--version"],
            capture_output=True,
            timeout=5,
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def _pandoc_convert(markdown: str, output_path: Path, to: str) -> None:
    import tempfile

    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".md", encoding="utf-8", delete=False
    ) as tmp:
        tmp.write(markdown)
        tmp_path = tmp.name

    cmd = [
        "pandoc", tmp_path,
        "-o", str(output_path),
        "--standalone",
        "--toc",
    ]
    if to == "pdf":
        cmd += ["--pdf-engine=xelatex"]

    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

    import os
    os.unlink(tmp_path)

    if result.returncode != 0:
        raise RuntimeError(
            f"pandoc failed (exit {result.returncode}):\n{result.stderr}"
        )
    _log.metric(f"pandoc_{to}", path=str(output_path))
